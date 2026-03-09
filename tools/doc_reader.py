"""
文档读取工具 - 支持 .docx 和 .doc 格式
"""
import sys
from pathlib import Path

def read_docx(file_path):
    """读取 .docx 文件"""
    from docx import Document
    
    doc = Document(file_path)
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # 读取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text_parts.append(row_text)
    
    return "\n".join(text_parts)

def read_doc(file_path):
    """
    读取旧版 .doc 文件
    需要安装 antiword 或使用其他方法
    """
    # 尝试使用 win32com (Windows + Word)
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(file_path)
        text = doc.Content.Text
        doc.Close()
        word.Quit()
        return text
    except Exception as e:
        return f"无法读取 .doc 文件: {e}\n建议: 将 .doc 转换为 .docx 格式，或安装 Microsoft Word"

def read_document(file_path):
    """自动检测并读取文档"""
    path = Path(file_path)
    
    if not path.exists():
        return f"文件不存在: {file_path}"
    
    suffix = path.suffix.lower()
    
    if suffix == '.docx':
        return read_docx(file_path)
    elif suffix == '.doc':
        return read_doc(file_path)
    else:
        return f"不支持的格式: {suffix}（支持 .doc 和 .docx）"

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python doc_reader.py <文档路径>")
        sys.exit(1)
    
    result = read_document(sys.argv[1])
    print(result)
